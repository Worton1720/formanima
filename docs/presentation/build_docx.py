#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Генератор Word-документа с текстом выступления (только то, что читать вслух).
Текст берётся из defense-speech.md (раздел «1. Текст выступления»),
чтобы документы не расходились.

Запуск:  uv run --with python-docx python build_docx.py
Результат: defense-speech.docx  (шрифт 14, полуторный интервал)
"""
import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(HERE, "defense-speech.md")
OUT = os.path.join(HERE, "defense-speech.docx")

# --- Вытащить абзацы речи из markdown (между "## 1." и "## 2.") ---
with open(SRC, encoding="utf-8") as f:
    md = f.read()

block = md.split("## 1.")[1].split("## 2.")[0]
paragraphs = [ln.strip() for ln in block.splitlines() if ln.strip().startswith("[Слайд")]

# Убрать пояснения-«справки» в круглых скобках — их вслух не читают.
# Формат: _(справка: ...)_  (markdown-курсив). Также чистим двойные пробелы.
def strip_notes(text):
    text = re.sub(r"\s*_\(справка:.*?\)_", "", text)
    text = re.sub(r"\s{2,}", " ", text)
    return text.strip()

paragraphs = [strip_notes(p) for p in paragraphs]

doc = Document()

# Базовый стиль: Times New Roman 14, полуторный интервал
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(14)
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
pf.space_after = Pt(6)

# Поля документа
sec = doc.sections[0]
sec.top_margin = Cm(2)
sec.bottom_margin = Cm(2)
sec.left_margin = Cm(3)
sec.right_margin = Cm(1.5)

# Заголовок
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Текст выступления на защите ВКР")
r.bold = True
r.font.size = Pt(14)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run("«Разработка веб-приложения для трекинга привычек "
                 "и целей с элементами геймификации»")
rs.italic = True
rs.font.size = Pt(14)

doc.add_paragraph()  # пустая строка

marker_re = re.compile(r"^(\[Слайд\s*\d+\])\s*(.*)$", re.S)
for para in paragraphs:
    m = marker_re.match(para)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    if m:
        mr = p.add_run(m.group(1) + " ")
        mr.bold = True
        mr.font.color.rgb = RGBColor(0x99, 0x66, 0x00)
        p.add_run(m.group(2))
    else:
        p.add_run(para)

doc.save(OUT)
print("OK saved:", OUT, "| paragraphs:", len(paragraphs))
